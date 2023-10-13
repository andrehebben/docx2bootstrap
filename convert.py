import subprocess
import os
from docx import Document
import datetime

class DOCXConverter:
    def __init__(self, input_directory, output_directory, template_file, css_file, lua_filters):
        self.input_directory = input_directory
        self.output_directory = output_directory
        self.template_file = template_file
        self.css_file = css_file
        self.lua_filters = lua_filters
        self.result_directory = output_directory

    def accept_all_changes(self, doc):
        for paragraph in doc.paragraphs:
            for run in paragraph.runs:
                run.clear()

    def convert_docx_to_html(self, file_name):
        file_path = os.path.join(self.input_directory, file_name)
        doc = Document(file_path)

        # Accept all changes in the source document
        self.accept_all_changes(doc)

        header_to_template = {
            "foundational": "./templates/chapter1.html",
            "cross sectional": "./templates/chapter2.html",
            "ecs key application areas": "./templates/chapter3.html",
            # Add more mappings as needed
        }

        selected_template = self.template_file

        for paragraph in doc.paragraphs:
            paragraph_text = paragraph.text
            paragraph_text_lower = paragraph_text.lower()

            template = header_to_template.get(paragraph_text_lower)

            if template:
                selected_template = template
                break

        output_html_file = os.path.join(self.result_directory, f"{datetime.datetime.now().strftime('%Y%m%d_%H%M%S')}_{file_name}.html")
        html_file_name = os.path.basename(output_html_file)
        pandoc_command = [
            "pandoc",
            file_path,
            "-o",
            output_html_file,
            "--template",
            selected_template,
            "--css",
            self.css_file,
            "--toc",
            "--toc-depth",
            "3",
        ]

        for lua_filter in self.lua_filters:
            pandoc_command.extend(["--lua-filter", lua_filter])

        pandoc_command.append("--extract-media=media")

        try:
            subprocess.run(pandoc_command, check=True)

            result_file_path = os.path.join(self.input_directory, f"conversion_result_{datetime.datetime.now().strftime('%Y%m%d_%H%M%S')}_{file_name}.txt")
            with open(result_file_path, "w") as result_file:
                result_file.write("CONVERSION SUCCEEDED\n")
                result_file.write("=" * 20 + "\n\n")
                result_file.write("The document was successfully converted:\n")
                result_file.write("-" * 20 + "\n")
                result_file.write(f"{file_name}\n")
                result_file.write("\n")
                result_file.write("The Sandbox website can be viewed here:\n")
                result_file.write("-" * 30 + "\n")
                result_file.write(f"https://sandbox.ecssria.eu/{html_file_name}\n")

            print(f"Conversion succeeded for {file_name}. Result written to '{result_file_path}'")
            os.remove(file_path)

        except subprocess.CalledProcessError as e:
            result_file_path = os.path.join(self.input_directory, f"{datetime.datetime.now().strftime('%Y%m%d_%H%M%S')}_{file_name}.txt")
            with open(result_file_path, "w") as result_file:
                result_file.write("CONVERSION FAILED\n")
                result_file.write("=" * 20 + "\n\n")
                result_file.write("Error details:\n")
                result_file.write(str(e) + "\n")

            print(f"Error: Pandoc conversion failed for {file_name} with error code {e.returncode}. Error details written to '{result_file_path}'")

    def convert_all_files(self):
        docx_files = [file for file in os.listdir(self.input_directory) if file.endswith(".docx")]
        sorted_docx_files = sorted(docx_files)

        if not docx_files:
            print("No DOCX files found in the input directory.")
        else:
            for file_name in sorted_docx_files:
                self.convert_docx_to_html(file_name)
            print("Cleanup completed: Removed input DOCX files.")

if __name__ == "__main__":
    input_directory = "./output"
    output_directory = "./web"
    template_file = "template.html"
    css_file = "template.css"
    lua_filters = [
        "filters/figure-captions.lua",
        "filters/accordion.lua",
        # Add more filters as needed
    ]

    converter = DOCXConverter(input_directory, output_directory, template_file, css_file, lua_filters)
    converter.convert_all_files()
