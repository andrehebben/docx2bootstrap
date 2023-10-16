import unittest
from unittest.mock import patch, call
from docx import Document
from convert import DOCXConverter

class TestDocxConverter(unittest.TestCase):
    @classmethod
    def setUpClass(cls):
        # Create a test directory and copy sample DOCX files for testing
        # You can also use temporary directories for this
        pass

    @classmethod
    def tearDownClass(cls):
        # Remove any test files or directories created during testing
        pass

    def setUp(self):
        # Initialize a DocxConverter instance for testing
        self.converter = DOCXConverter(input_directory="./test_input", output_directory="./test_output", template_file="template.html", css_file="template.css", lua_filters=["filters/figure-captions.lua"])

    def tearDown(self):
        # Clean up any test-specific resources
        pass

    @patch('builtins.open', new_callable=unittest.mock.mock_open)
    def test_accept_all_changes(self, mock_open):
        # Create a test Document
        doc = Document()
        doc.add_paragraph("Test paragraph")
        run = doc.paragraphs[0].runs[0]
        run.add_text("Text to clear")

        # Accept changes
        self.converter.accept_all_changes(doc)

        # Assert that the run is cleared
        self.assertEqual(run.text, '')

    @patch('subprocess.run')
    @patch('os.makedirs')
    def test_convert_docx_to_html(self, mock_makedirs, mock_subprocess_run):
        # Prepare test data
        file_name = "test_document.docx"
        header_to_template = {
            "test header": "test_template.html"
        }

        # Mock subprocess.run to avoid executing pandoc during testing
        mock_subprocess_run.return_value.returncode = 0

        # Run the conversion method
        self.converter.convert_docx_to_html(file_name, header_to_template)

        # Check that subprocess.run is called with the correct arguments
        pandoc_command = [
            "pandoc",
            f"./test_input/{file_name}",
            "-o",
            f"./test_output/result/{file_name}.html",
            "--template",
            "test_template.html",
            "--css",
            "template.css",
            "--toc",
            "--toc-depth",
            "3",
            "--lua-filter",
            'filters/figure-captions.lua',
            '--extract-media=media',
        ]
        mock_subprocess_run.assert_called_once_with(pandoc_command, check=True)

    @patch('os.listdir')
    @patch('os.path.join')
    @patch('os.remove')
    @patch('os.path.basename')
    @patch('subprocess.run')
    @patch('builtins.open', new_callable=unittest.mock.mock_open)
    @patch('os.makedirs')
    def test_process_files(self, mock_makedirs, mock_open, mock_subprocess_run, mock_basename, mock_remove, mock_join, mock_listdir):
        # Mock the list of DOCX files in the input directory
        mock_listdir.return_value = ["test_document1.docx", "test_document2.docx"]

        # Mock the behavior of the open function here
        mock_open.side_effect = lambda file_path, mode: file_path

        # Mock various functions to avoid actual execution during testing
        mock_join.side_effect = lambda dir, file: f"{dir}/{file}"
        mock_basename.side_effect = lambda path: path.split('/')[-1
        mock_subprocess_run.return_value.returncode = 0

        # Run the process_files method
        header_to_template = {
            "test header": "test_template.html"
        }
        self.converter.process_files(header_to_template)

        # Check that subprocess.run is called with the correct arguments
        expected_pandoc_command = [
            call(['pandoc', './test_input/test_document1.docx', '-o', './test_output/result/test_document1.docx.html', '--template', 'test_template.html', '--css', 'template.css', '--toc', '--toc-depth', '3', '--lua-filter', 'filters/figure-captions.lua', '--extract-media=media']),
            call(['pandoc', './test_input/test_document2.docx', '-o', './test_output/result/test_document2.docx.html', '--template', 'test_template.html', '--css', 'template.css', '--toc', '--toc-depth', '3', '--lua-filter', 'filters/figure-captions.lua', '--extract-media=media']),
        ]
        mock_subprocess_run.assert_has_calls(expected_pandoc_command)

        # Check that the 'open' function was called to write result files
        self.assertEqual(mock_open.call_count, 2)

        # Add more tests as needed for your specific use case

        if __name__ == "__main__":
            unittest.main()
